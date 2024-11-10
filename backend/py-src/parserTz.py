from docx import Document
import re

def clean_text(text):
    """
    Убирает лишние символы, такие как неразрывные пробелы (\xa0) и переводы строк (\n).
    """
    text = text.replace('\xa0', ' ')  # Заменяем неразрывный пробел на обычный
    text = text.replace('\n', ' ')  # Убираем переводы строк
    text = re.sub(r'\s+', ' ', text.strip())  # Убираем лишние пробелы
    return text.strip()

def parse_characteristics(characteristics_text):
    """
    Разбивает строку характеристик на ключ-значение.
    """
    characteristics = {}
    pairs = re.split(r'(?<!:)\s*([А-Я][^:]*):', characteristics_text)  # Разбивает по шаблону "Ключ: значение"
    for i in range(1, len(pairs), 2):  # Пропускаем пустые части
        key = clean_text(pairs[i])
        value = clean_text(pairs[i + 1]) if i + 1 < len(pairs) else ''
        characteristics[key] = value
    return characteristics

def parse_tz_docx(file_path):
    """
    Функция для обработки ТЗ в формате DOCX и извлечения необходимых данных.
    """
    doc = Document(file_path)
    data = {
        "contract": 0,  # По умолчанию "обеспечение контракта" отсутствует
        "law": "44-ФЗ",  # Закон, согласно которому происходит заключение
        "deadlines": None,  # Сроки поставки (если указаны)
        "products": []  # Список товаров
    }

    # Проверка общего текста документа
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Проверка на наличие обязательства обеспечения контракта
        if "обеспечение исполнения контракта" in text.lower():
            data["contract"] = 1

        # Извлечение информации о сроках поставки
        deadlines_match = re.search(r"срок поставки.*?(\d+)\s*(дней|дня|суток)", text.lower())
        if deadlines_match:
            data["deadlines"] = deadlines_match.group(1)

    # Парсинг таблиц для извлечения товаров и их характеристик
    for table in doc.tables:
        headers = []  # Список заголовков таблицы (первой строки)
        for i, row in enumerate(table.rows):
            cells = [clean_text(cell.text) for cell in row.cells]

            # Если это первая строка, предполагаем, что это заголовки
            if i == 0:
                headers = [header.lower() for header in cells]  # Приводим заголовки к нижнему регистру
                continue

            # Иначе обрабатываем как строки данных
            if len(cells) > 1:  # Игнорируем пустые строки
                product = {}
                for j, cell in enumerate(cells):
                    if j < len(headers):  # Убедимся, что индекс в рамках заголовков
                        key = headers[j]
                        product[key] = cell

                # Преобразуем каждую характеристику в ключ-значение
                if "наименование" in product and "кол-во" in product:
                    formatted_product = {
                        "name": product.get("наименование", "Не указано"),
                        "quantity": product.get("кол-во", "0")
                    }

                    # Разбиваем характеристики на отдельные ключи, если они есть
                    if "характеристика" in product:
                        characteristics = parse_characteristics(product["характеристика"])
                        formatted_product.update(characteristics)

                    # Добавляем остальные параметры
                    for key, value in product.items():
                        if key not in ["наименование", "кол-во", "характеристика"]:
                            formatted_product[key] = value

                    data["products"].append(formatted_product)

    return data